from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from typing import Optional, List
import os
import json
import io
from datetime import datetime
from dotenv import load_dotenv
import google.generativeai as genai
from gtts import gTTS
import PyPDF2
from docx import Document
from fpdf import FPDF
import markdown
from difflib import SequenceMatcher
import tempfile

# Load environment variables
load_dotenv()

app = FastAPI(title="TextMorph API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    print("WARNING: GEMINI_API_KEY not found in environment variables")
    print("The API will start but text processing features will not work")
    model = None
else:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        print("Gemini API initialized successfully")
    except Exception as e:
        print(f"WARNING: Failed to initialize Gemini API: {e}")
        model = None

# In-memory storage (replace with database in production)
users_db = {}
user_profiles = {}
user_history = {}
user_sessions = {}

# Pydantic models
class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    password: str

class ProcessRequest(BaseModel):
    operation: str  # "summarize", "paraphrase", "translate_summarize"
    text: str
    language: str = "English"
    tone: str = "Professional"
    adaptation: str = "General"
    summary_type: str = "Abstractive"
    depth_level: str = "Detailed"
    readability_level: str = "General public"
    content_type: str = "auto"

class CompareRequest(BaseModel):
    text1: str
    text2: str

class TextToSpeechRequest(BaseModel):
    text: str
    language: str = "en"

class ExportRequest(BaseModel):
    text: str
    title: str = "TextMorph_Export"
    format: str  # "pdf", "docx", "md", "html", "json"

class HistoryAddRequest(BaseModel):
    username: str
    operation: str
    params: dict
    output: str

class ProfileUpdateRequest(BaseModel):
    username: str
    tone: Optional[str] = None
    style: Optional[str] = None
    language: Optional[str] = None
    depth: Optional[str] = None

class ClearHistoryRequest(BaseModel):
    username: str

# Helper functions
def detect_content_type(text: str) -> str:
    """Detect content type from text"""
    text_lower = text.lower()
    if any(word in text_lower for word in ['abstract', 'introduction', 'methodology', 'conclusion']):
        return "Academic"
    elif any(word in text_lower for word in ['code', 'function', 'variable', 'api']):
        return "Technical"
    elif any(word in text_lower for word in ['marketing', 'campaign', 'customer', 'product']):
        return "Business"
    elif any(word in text_lower for word in ['news', 'report', 'article', 'journalism']):
        return "Journalistic"
    else:
        return "General"

def generate_prompt(operation: str, text: str, language: str, tone: str, adaptation: str,
                   summary_type: str, depth_level: str, readability_level: str, content_type: str) -> str:
    """Generate prompt for Gemini based on operation and parameters"""
    
    if content_type == "auto":
        detected_type = detect_content_type(text)
    else:
        detected_type = content_type
    
    if operation == "summarize":
        depth_map = {
            "Brief": "in 2-3 sentences",
            "Detailed": "in 5-7 sentences",
            "Comprehensive": "in 10+ sentences with key points"
        }
        depth_instruction = depth_map.get(depth_level, "in detail")
        
        prompt = f"""Summarize the following {detected_type.lower()} text in {language} with a {tone.lower()} tone.
        
        Requirements:
        - Summary type: {summary_type}
        - Depth: {depth_instruction}
        - Readability: {readability_level}
        - Adaptation: {adaptation}
        
        Text to summarize:
        {text}
        
        Provide a clear, well-structured summary."""
        
    elif operation == "paraphrase":
        prompt = f"""Paraphrase the following text in {language} with a {tone.lower()} tone.
        
        Requirements:
        - Maintain the original meaning
        - Use different wording and sentence structure
        - Adaptation style: {adaptation}
        - Readability: {readability_level}
        
        Text to paraphrase:
        {text}
        
        Provide a paraphrased version."""
        
    elif operation == "translate_summarize":
        prompt = f"""Translate and summarize the following text to {language} with a {tone.lower()} tone.
        
        Requirements:
        - First translate to {language}
        - Then summarize {depth_level.lower()}
        - Summary type: {summary_type}
        - Readability: {readability_level}
        
        Text:
        {text}
        
        Provide the translated summary."""
    
    else:
        prompt = f"Process the following text: {text}"
    
    return prompt

# Authentication endpoints
@app.post("/api/login")
async def login(request: LoginRequest):
    if request.username in users_db:
        if users_db[request.username] == request.password:
            # Initialize profile if new
            if request.username not in user_profiles:
                user_profiles[request.username] = {
                    "tone": "Professional",
                    "style": "General",
                    "language": "English",
                    "depth": "Detailed"
                }
            if request.username not in user_history:
                user_history[request.username] = []
            
            return {
                "success": True,
                "profile": user_profiles[request.username]
            }
        else:
            raise HTTPException(status_code=401, detail="Invalid password")
    else:
        raise HTTPException(status_code=404, detail="User not found")

@app.post("/api/register")
async def register(request: RegisterRequest):
    if request.username in users_db:
        raise HTTPException(status_code=400, detail="Username already exists")
    
    users_db[request.username] = request.password
    user_profiles[request.username] = {
        "tone": "Professional",
        "style": "General",
        "language": "English",
        "depth": "Detailed"
    }
    user_history[request.username] = []
    
    return {"success": True, "message": "User registered successfully"}

@app.post("/api/logout")
async def logout():
    return {"success": True, "message": "Logged out successfully"}

# Health check endpoint
@app.get("/api/health")
async def health_check():
    return {
        "status": "ok",
        "gemini_configured": model is not None,
        "timestamp": datetime.now().isoformat()
    }

# Text extraction endpoints
@app.post("/api/extract/pdf")
async def extract_pdf(file: UploadFile = File(...)):
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    try:
        contents = await file.read()
        pdf_file = io.BytesIO(contents)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return {"text": text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting PDF: {str(e)}")

@app.post("/api/extract/docx")
async def extract_docx(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="File must be a DOCX file")
    
    try:
        contents = await file.read()
        doc_file = io.BytesIO(contents)
        doc = Document(doc_file)
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {"text": text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting DOCX: {str(e)}")

# AI processing endpoint
@app.post("/api/process")
async def process_text(request: ProcessRequest):
    if model is None:
        raise HTTPException(status_code=503, detail="Gemini API is not configured. Please set GEMINI_API_KEY in .env file")
    
    try:
        prompt = generate_prompt(
            request.operation,
            request.text,
            request.language,
            request.tone,
            request.adaptation,
            request.summary_type,
            request.depth_level,
            request.readability_level,
            request.content_type
        )
        
        response = model.generate_content(prompt)
        result = response.text
        
        return {"result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing text: {str(e)}")

# Text comparison endpoint
@app.post("/api/compare")
async def compare_texts(request: CompareRequest):
    if model is None:
        raise HTTPException(status_code=503, detail="Gemini API is not configured. Please set GEMINI_API_KEY in .env file")
    
    try:
        # Basic similarity using difflib
        similarity_ratio = SequenceMatcher(None, request.text1, request.text2).ratio()
        basic_similarity = f"{similarity_ratio * 100:.1f}%"
        
        # AI analysis using Gemini
        comparison_prompt = f"""Compare these two texts and provide a detailed analysis:
        
        Text 1:
        {request.text1}
        
        Text 2:
        {request.text2}
        
        Analyze:
        1. Similarities in content and meaning
        2. Differences in tone, style, and structure
        3. Key distinctions
        4. Overall assessment"""
        
        response = model.generate_content(comparison_prompt)
        ai_analysis = response.text
        
        return {
            "ai_analysis": ai_analysis,
            "basic_similarity": basic_similarity
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error comparing texts: {str(e)}")

# Text-to-speech endpoint
@app.post("/api/text-to-speech")
async def text_to_speech(request: TextToSpeechRequest):
    try:
        # Language code mapping
        lang_map = {
            "English": "en",
            "Spanish": "es",
            "French": "fr",
            "German": "de",
            "Chinese": "zh",
            "Japanese": "ja"
        }
        lang_code = lang_map.get(request.language, request.language)
        
        tts = gTTS(text=request.text, lang=lang_code, slow=False)
        audio_buffer = io.BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(audio_buffer.read()),
            media_type="audio/mpeg",
            headers={"Content-Disposition": "attachment; filename=output.mp3"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating speech: {str(e)}")

# Export endpoints
@app.post("/api/export")
async def export_text(request: ExportRequest):
    try:
        if request.format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Handle text wrapping
            lines = request.text.split('\n')
            for line in lines:
                pdf.cell(0, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
            
            output = io.BytesIO()
            pdf.output(output)
            output.seek(0)
            
            return StreamingResponse(
                io.BytesIO(output.read()),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={request.title}.pdf"}
            )
        
        elif request.format == "docx":
            doc = Document()
            doc.add_paragraph(request.text)
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return StreamingResponse(
                io.BytesIO(output.read()),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={request.title}.docx"}
            )
        
        elif request.format == "md":
            return StreamingResponse(
                io.BytesIO(request.text.encode('utf-8')),
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename={request.title}.md"}
            )
        
        elif request.format == "html":
            html_content = markdown.markdown(request.text)
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{request.title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            return StreamingResponse(
                io.BytesIO(full_html.encode('utf-8')),
                media_type="text/html",
                headers={"Content-Disposition": f"attachment; filename={request.title}.html"}
            )
        
        elif request.format == "json":
            json_data = {
                "title": request.title,
                "content": request.text,
                "exported_at": datetime.now().isoformat()
            }
            
            return StreamingResponse(
                io.BytesIO(json.dumps(json_data, indent=2).encode('utf-8')),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename={request.title}.json"}
            )
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported export format")
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting file: {str(e)}")

# History endpoints
@app.get("/api/history")
async def get_history(username: str):
    if username not in user_history:
        return []
    
    history = user_history[username]
    # Format for frontend
    formatted_history = []
    for item in history:
        formatted_history.append({
            "id": item.get("id", ""),
            "operation": item.get("operation", ""),
            "timestamp": item.get("timestamp", ""),
            "preview": item.get("output", "")[:100] + "..." if len(item.get("output", "")) > 100 else item.get("output", ""),
            "fullOutput": item.get("output", ""),
            "parameters": item.get("params", {})
        })
    
    return formatted_history

@app.post("/api/history/add")
async def add_history(request: HistoryAddRequest):
    if request.username not in user_history:
        user_history[request.username] = []
    
    history_item = {
        "id": str(len(user_history[request.username]) + 1),
        "operation": request.operation,
        "timestamp": datetime.now().isoformat(),
        "params": request.params,
        "output": request.output
    }
    
    user_history[request.username].append(history_item)
    return {"success": True, "message": "History item added"}

@app.post("/api/history/clear")
async def clear_history(request: ClearHistoryRequest):
    if request.username in user_history:
        user_history[request.username] = []
    return {"success": True, "message": "History cleared"}

# Profile endpoints
@app.get("/api/profile")
async def get_profile(username: str):
    if username not in user_profiles:
        raise HTTPException(status_code=404, detail="User profile not found")
    
    profile = user_profiles[username]
    return {
        "preferred_tone": profile.get("tone", "Professional"),
        "preferred_style": profile.get("style", "General"),
        "preferred_language": profile.get("language", "English"),
        "preferred_depth": profile.get("depth", "Detailed")
    }

@app.post("/api/profile/update")
async def update_profile(request: ProfileUpdateRequest):
    if request.username not in user_profiles:
        user_profiles[request.username] = {}
    
    if request.tone:
        user_profiles[request.username]["tone"] = request.tone
    if request.style:
        user_profiles[request.username]["style"] = request.style
    if request.language:
        user_profiles[request.username]["language"] = request.language
    if request.depth:
        user_profiles[request.username]["depth"] = request.depth
    
    return {"success": True, "message": "Profile updated", "profile": user_profiles[request.username]}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

